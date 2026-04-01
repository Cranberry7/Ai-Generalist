import os
import json
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

def generate_docx_report(phase3_json_path: str, output_docx_path: str) -> bool:
    """Read the vetted Phase 3 JSON payload and generate a formatted DOCX."""
    if not os.path.exists(phase3_json_path):
        logger.error(f"Cannot find Phase 3 payload at {phase3_json_path}")
        return False
        
    try:
        with open(phase3_json_path, 'r', encoding='utf-8') as f:
            diagnostic_data = json.load(f)
            
        logger.info("Initializing Document Assembler...")
        doc = Document()
        
        # 1. Document Title
        title = doc.add_heading('Detailed Diagnostic Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. Project Metadata
        project_info = diagnostic_data.get("project_info", {})
        doc.add_heading('Project Information', level=1)
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = "Client Name"
        table.cell(0, 1).text = project_info.get("client_name", "N/A")
        
        table.cell(1, 0).text = "Inspection Date"
        table.cell(1, 1).text = project_info.get("date", "N/A")
        
        table.cell(2, 0).text = "Property Address"
        table.cell(2, 1).text = project_info.get("address", "N/A")
        
        doc.add_page_break()
        
        # 3. Observations
        doc.add_heading('Inspection Findings by Area', level=1)
        observations = diagnostic_data.get("observations", [])
        
        for obs in observations:
            area_name = obs.get("area_name", "Unknown Area")
            doc.add_heading(area_name, level=2)
            
            # Diagnostic Statement
            statement = obs.get("unified_diagnostic_statement", "")
            doc.add_paragraph(statement)
            
            # Conflict Alert
            if obs.get("conflict_detected"):
                conflict_p = doc.add_paragraph()
                run = conflict_p.add_run("⚠️ HEURISTIC WARNING: " + obs.get("conflict_reason", ""))
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0) # Red
            
            # Images
            images = obs.get("associated_images", [])
            for img in images:
                abs_path = img.get("absolute_path")
                if abs_path and os.path.exists(abs_path):
                    try:
                        # Insert image securely fitted within margins (e.g. 5 inches wide)
                        doc.add_picture(abs_path, width=Inches(5.0))
                        
                        # Apply caption directly underneath
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        desc_run = caption.add_run(img.get("description", "No description provided."))
                        desc_run.italic = True
                        desc_run.font.size = Pt(9)
                        
                        # Add relevance
                        doc.add_paragraph("Relevance: " + img.get("relevance", ""))
                    except Exception as img_err:
                        logger.warning(f"Failed to insert image {abs_path}: {img_err}")
                        
            doc.add_paragraph() # Spacing between areas

        # 4. Missing Information
        missing_info = diagnostic_data.get("missing_info", [])
        if missing_info:
            doc.add_page_break()
            doc.add_heading('Missing Document Information', level=1)
            for info in missing_info:
                doc.add_paragraph(info, style='List Bullet')

        # 5. Export
        doc.save(output_docx_path)
        logger.info(f"Phase 4 Assembly complete! Document compiled to {output_docx_path}")
        return True
        
    except Exception as e:
        logger.error(f"Assembly Failure: {str(e)}")
        return False
